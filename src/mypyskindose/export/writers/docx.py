"""DOCX report writer (§6) via ``python-docx`` (optional ``export`` extra).

Mirrors the PDF section order. ``render_docx_bytes(payload)`` for downloads;
``write_docx(payload, path)`` for filesystem writes.
"""

from __future__ import annotations

import io
from pathlib import Path

from docx import Document  # type: ignore[import-untyped]
from docx.shared import Inches, Pt, RGBColor  # type: ignore[import-untyped]

from .._format import (
    CORRECTION_HEADER,
    OFFSET_LABELS,
    collect_alert_lines,
    correction_row,
    dosimetric_rows,
)
from ..models import ExportPayload

_AMBER = RGBColor(0xFF, 0xF3, 0xCD)
_RED = RGBColor(0xF8, 0xD7, 0xDA)


def _shade(cell, color: RGBColor) -> None:
    """Apply a solid background fill to a table cell (no direct API in docx)."""
    from docx.oxml import OxmlElement  # type: ignore[import-untyped]
    from docx.oxml.ns import qn  # type: ignore[import-untyped]

    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), f"{color}")
    cell._tc.get_or_add_tcPr().append(shd)


def _table(doc, rows: list[list], *, header: bool = True):
    t = doc.add_table(rows=0, cols=len(rows[0]))
    t.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        for c, value in enumerate(row):
            cells[c].text = str(value)
            if header and i == 0:
                for p in cells[c].paragraphs:
                    for run in p.runs:
                        run.font.bold = True
    return t


def _alerts(doc, payload: ExportPayload) -> None:
    doc.add_heading("Executive alerts", level=2)
    for text, severity in collect_alert_lines(payload):
        t = doc.add_table(rows=1, cols=1)
        cell = t.rows[0].cells[0]
        cell.text = text
        if severity == "error":
            _shade(cell, _RED)
        elif severity == "warning":
            _shade(cell, _AMBER)


def _settings_rows(exam) -> list[list[str]]:
    rows = [["Setting", "Value"]]
    for key, value in exam.settings.items():
        if key in ("phantom", "patient_offset"):
            continue
        rows.append([key, str(value)])
    for key, value in exam.settings.get("phantom", {}).items():
        rows.append([f"phantom.{key}", str(value)])
    for key, value in exam.settings.get("patient_offset", {}).items():
        rows.append([str(OFFSET_LABELS.get(key, key)), str(value)])
    rows.append(["Manufacturer", exam.manufacturer or "N/A"])
    rows.append(["Model", exam.model or "N/A"])
    return rows


def build_document(payload: ExportPayload):
    doc = Document()
    doc.add_heading(payload.meta.report_title, level=0)
    p = doc.add_paragraph()
    run = p.add_run(
        f"{payload.meta.app_name} v{payload.meta.package_version} · "
        f"{payload.meta.execution_context} · {payload.meta.generated_at.isoformat(timespec='seconds')}"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    _alerts(doc, payload)

    doc.add_heading("Cumulative summary", level=2)
    _table(doc, [["Metric", "Value"]] + dosimetric_rows(payload.cumulative.metrics))

    if payload.is_multi_exam:
        doc.add_heading("Per-exam results", level=2)
        headers = ["Metric"] + [e.exam_id for e in payload.exams] + ["Cumulative"]
        names = [r[0] for r in dosimetric_rows(payload.cumulative.metrics)]
        cols = [{r[0]: r[1] for r in dosimetric_rows(e.metrics)} for e in payload.exams]
        cum = {r[0]: r[1] for r in dosimetric_rows(payload.cumulative.metrics)}
        summary = [headers] + [[n, *[c.get(n, "N/A") for c in cols], cum.get(n, "N/A")] for n in names]
        _table(doc, summary)

    doc.add_heading("Settings & equipment", level=2)
    for exam in payload.exams:
        if payload.is_multi_exam:
            doc.add_heading(f"Exam {exam.exam_id}", level=3)
        _table(doc, _settings_rows(exam))

    doc.add_heading("Correction factors", level=2)
    if payload.is_multi_exam:
        for exam in payload.exams:
            doc.add_paragraph(f"Exam {exam.exam_id}")
            _table(doc, [CORRECTION_HEADER] + [correction_row(s) for s in exam.corrections])
        doc.add_paragraph("Cumulative (kerma-weighted)")
    _table(doc, [CORRECTION_HEADER] + [correction_row(s) for s in payload.cumulative.corrections])

    if payload.images:
        doc.add_heading("Dose-map images", level=2)
        for entry in payload.images:
            doc.add_paragraph(entry.label)
            if entry.png_bytes is not None:
                doc.add_picture(io.BytesIO(entry.png_bytes), width=Inches(6.0))
            else:
                doc.add_paragraph(entry.error_message or "Image unavailable (kaleido/export error)")
    return doc


def render_docx_bytes(payload: ExportPayload) -> bytes:
    buf = io.BytesIO()
    build_document(payload).save(buf)
    return buf.getvalue()


def write_docx(payload: ExportPayload, path: Path) -> None:
    build_document(payload).save(str(Path(path)))
