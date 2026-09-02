"""DOCX report writer (§6) via ``python-docx`` (optional ``export`` extra).

Mirrors the PDF section order. ``render_docx_bytes(payload)`` for downloads;
``write_docx(payload, path)`` for filesystem writes.
"""

from __future__ import annotations

import io
from pathlib import Path

from docx import Document  # type: ignore[import-untyped]
from docx.shared import Inches, Pt, RGBColor  # type: ignore[import-untyped]

from guiskindose.safe_output import atomic_write_private

from .._format import (
    CORRECTION_HEADER,
    KERMA_METER_WEIGHTING_FOOTNOTE,
    OFFSET_LABELS,
    collect_alert_lines,
    correction_row,
    corrections_use_kerma_meter,
    dosimetric_rows,
)
from ..models import ExportPayload

_AMBER = RGBColor(0xFF, 0xF3, 0xCD)
_RED = RGBColor(0xF8, 0xD7, 0xDA)


def _shade(cell, color: RGBColor) -> None:
    """Apply a solid background fill to a table cell (no direct API in docx)."""
    from docx.oxml import OxmlElement  # type: ignore[import-untyped]
    from docx.oxml.ns import qn  # type: ignore[import-untyped]

    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), f"{color}")
    cell._tc.get_or_add_tcPr().append(shd)


def _table(doc, rows: list[list], *, header: bool = True):
    """Append a simple DOCX table from row data."""
    t = doc.add_table(rows=0, cols=len(rows[0]))
    t.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        _add_table_row(t, row, bold=header and i == 0)
    return t


def _add_table_row(table, row: list, *, bold: bool) -> None:
    """Add one DOCX table row and optionally emphasize every cell value."""
    for cell, value in zip(table.add_row().cells, row, strict=True):
        cell.text = str(value)
        if bold:
            _bold_cell_runs(cell)


def _bold_cell_runs(cell) -> None:
    """Apply bold styling to the runs generated for a header cell."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True


def _alerts(doc, payload: ExportPayload) -> None:
    """Add the executive-alerts section to a DOCX document."""
    doc.add_heading("Executive alerts", level=2)
    for text, severity in collect_alert_lines(payload):
        t = doc.add_table(rows=1, cols=1)
        cell = t.rows[0].cells[0]
        cell.text = text
        if severity == "error":
            _shade(cell, _RED)
        elif severity == "warning":
            _shade(cell, _AMBER)


def _settings_rows(exam) -> list[list[str]]:
    """Build Setting/Value rows for one exam's settings block."""
    rows = [["Setting", "Value"]]
    for key, value in exam.settings.items():
        if key in ("phantom", "patient_offset"):
            continue
        rows.append([key, str(value)])
    for key, value in exam.settings.get("phantom", {}).items():
        rows.append([f"phantom.{key}", str(value)])
    for key, value in exam.settings.get("patient_offset", {}).items():
        rows.append([str(OFFSET_LABELS.get(key, key)), str(value)])
    rows.append(["Manufacturer", exam.manufacturer or "N/A"])
    rows.append(["Model", exam.model or "N/A"])
    for field_name, desc in exam.unit_conversions.items():
        rows.append([f"Units: {field_name}", desc])
    return rows


def build_document(payload: ExportPayload):
    """Assemble the full rich-export DOCX document."""
    doc = Document()
    _add_document_header(doc, payload)
    _alerts(doc, payload)
    _add_result_sections(doc, payload)
    _add_settings_section(doc, payload)
    _add_corrections_section(doc, payload)
    _add_image_section(doc, payload)
    return doc


def _add_document_header(doc, payload: ExportPayload) -> None:
    """Add the report title and compact provenance line."""
    doc.add_heading(payload.meta.report_title, level=0)
    p = doc.add_paragraph()
    run = p.add_run(
        f"{payload.meta.app_name} v{payload.meta.package_version} · "
        f"{payload.meta.execution_context} · {payload.meta.generated_at.isoformat(timespec='seconds')}"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)


def _add_result_sections(doc, payload: ExportPayload) -> None:
    """Add cumulative and, when applicable, per-exam dosimetric tables."""
    doc.add_heading("Cumulative summary", level=2)
    _table(doc, [["Metric", "Value"], *dosimetric_rows(payload.cumulative.metrics)])
    if payload.is_multi_exam:
        _add_per_exam_results(doc, payload)


def _add_per_exam_results(doc, payload: ExportPayload) -> None:
    """Add the multi-exam comparison table in stable metric order."""
    doc.add_heading("Per-exam results", level=2)
    headers = ["Metric"] + [e.exam_id for e in payload.exams] + ["Cumulative"]
    names = [row[0] for row in dosimetric_rows(payload.cumulative.metrics)]
    exam_columns = [{row[0]: row[1] for row in dosimetric_rows(exam.metrics)} for exam in payload.exams]
    cumulative = {row[0]: row[1] for row in dosimetric_rows(payload.cumulative.metrics)}
    rows = [[name, *[column.get(name, "N/A") for column in exam_columns], cumulative.get(name, "N/A")] for name in names]
    _table(doc, [headers, *rows])


def _add_settings_section(doc, payload: ExportPayload) -> None:
    """Add one settings table per exam, retaining multi-exam headings."""
    doc.add_heading("Settings & equipment", level=2)
    for exam in payload.exams:
        if payload.is_multi_exam:
            doc.add_heading(f"Exam {exam.exam_id}", level=3)
        _table(doc, _settings_rows(exam))


def _add_corrections_section(doc, payload: ExportPayload) -> None:
    """Add correction-factor tables and optional kerma-meter footnote."""
    doc.add_heading("Correction factors", level=2)
    if payload.is_multi_exam:
        for exam in payload.exams:
            doc.add_paragraph(f"Exam {exam.exam_id}")
            _table(doc, [CORRECTION_HEADER] + [correction_row(s) for s in exam.corrections])
        doc.add_paragraph("Cumulative (kerma-weighted)")
    _table(doc, [CORRECTION_HEADER] + [correction_row(s) for s in payload.cumulative.corrections])
    if corrections_use_kerma_meter(payload):
        doc.add_paragraph(KERMA_METER_WEIGHTING_FOOTNOTE)


def _add_image_section(doc, payload: ExportPayload) -> None:
    """Add available images or their safe error notices."""
    if payload.images:
        doc.add_heading("Dose-map images", level=2)
        for entry in payload.images:
            doc.add_paragraph(entry.label)
            if entry.png_bytes is not None:
                doc.add_picture(io.BytesIO(entry.png_bytes), width=Inches(6.0))
            else:
                doc.add_paragraph(entry.error_message or "Image unavailable (kaleido/export error)")


def render_docx_bytes(payload: ExportPayload) -> bytes:
    """Render the DOCX report to an in-memory bytes payload."""
    buf = io.BytesIO()
    build_document(payload).save(buf)
    return buf.getvalue()


def write_docx(payload: ExportPayload, path: Path) -> None:
    """Atomically write the DOCX report to *path*."""
    atomic_write_private(path, render_docx_bytes(payload))
